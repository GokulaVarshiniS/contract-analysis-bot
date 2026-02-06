# src/document_parser.py

import os
import pdfplumber
from docx import Document
import re
from typing import Dict, Optional
from langdetect import detect
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocumentParser:
    """Handles parsing of PDF, DOCX, and TXT files"""
    
    SUPPORTED_FORMATS = ['.pdf', '.docx', '.doc', '.txt']
    
    def __init__(self):
        self.raw_text = ""
        self.metadata = {}
        self.language = "en"
    
    def parse(self, file_path: str) -> Dict:
        """
        Main parsing method that routes to appropriate parser
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext not in self.SUPPORTED_FORMATS:
            raise ValueError(f"Unsupported file format: {ext}")
        
        if ext == '.pdf':
            self.raw_text = self._parse_pdf(file_path)
        elif ext in ['.docx', '.doc']:
            self.raw_text = self._parse_docx(file_path)
        elif ext == '.txt':
            self.raw_text = self._parse_txt(file_path)
        
        # Detect language
        self.language = self._detect_language(self.raw_text)
        
        # Extract metadata
        self.metadata = {
            'file_name': os.path.basename(file_path),
            'file_size': os.path.getsize(file_path),
            'file_type': ext,
            'language': self.language,
            'word_count': len(self.raw_text.split()),
            'character_count': len(self.raw_text)
        }
        
        return {
            'text': self.raw_text,
            'metadata': self.metadata,
            'language': self.language
        }
    
    def _parse_pdf(self, file_path: str) -> str:
        """Extract text from PDF files"""
        text = ""
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n\n"
        except Exception as e:
            logger.error(f"Error parsing PDF: {e}")
            raise
        
        return self._clean_text(text)
    
    def _parse_docx(self, file_path: str) -> str:
        """Extract text from DOCX files"""
        text = ""
        try:
            doc = Document(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
        except Exception as e:
            logger.error(f"Error parsing DOCX: {e}")
            raise
        
        return self._clean_text(text)
    
    def _parse_txt(self, file_path: str) -> str:
        """Extract text from TXT files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
        except UnicodeDecodeError:
            with open(file_path, 'r', encoding='latin-1') as f:
                text = f.read()
        
        return self._clean_text(text)
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        # Remove special characters but keep punctuation
        text = re.sub(r'[^\w\s.,;:!?()\'"-]', '', text)
        # Normalize quotes
        text = text.replace('"', '"').replace('"', '"')
        text = text.replace(''', "'").replace(''', "'")
        
        return text.strip()
    
    def _detect_language(self, text: str) -> str:
        """Detect if text is in English or Hindi"""
        try:
            lang = detect(text[:1000])  # Use first 1000 chars for detection
            if lang == 'hi':
                return 'hi'
            return 'en'
        except:
            return 'en'


class HindiTextNormalizer:
    """Handles Hindi to English normalization for NLP tasks"""
    
    def __init__(self):
        # Common legal terms in Hindi with English equivalents
        self.legal_terms = {
            'अनुबंध': 'contract',
            'समझौता': 'agreement',
            'पक्ष': 'party',
            'शर्तें': 'terms',
            'नियम': 'conditions',
            'दायित्व': 'liability',
            'क्षतिपूर्ति': 'indemnity',
            'समाप्ति': 'termination',
            'अधिकार क्षेत्र': 'jurisdiction',
            'मध्यस्थता': 'arbitration',
            'गोपनीयता': 'confidentiality',
            'बौद्धिक संपदा': 'intellectual property',
            'भुगतान': 'payment',
            'अवधि': 'duration',
            'नवीनीकरण': 'renewal'
        }
    
    def normalize(self, hindi_text: str) -> str:
        """
        Normalize Hindi text by replacing common legal terms
        Note: For full translation, integrate with translation API
        """
        normalized = hindi_text
        for hindi, english in self.legal_terms.items():
            normalized = normalized.replace(hindi, english)
        return normalized